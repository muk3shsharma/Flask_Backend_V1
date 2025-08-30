"""
Document Processing Utilities Module - API Version
=================================================

Core utilities for Word document manipulation optimized for API use.
Handles text replacement, image insertion, and document structure manipulation.
"""
import os
import docx
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from werkzeug.utils import secure_filename
import time
import tempfile


def find_and_replace_text(doc, old_text, new_text):
    """
    Robust text replacement that handles Word's run structure.
    Specifically handles cases where placeholders are split across runs.
    """
    def replace_in_paragraph(paragraph):
        """Replace text in a single paragraph, handling split runs."""
        if not paragraph.runs:
            return False
        
        # Method 1: Check if any single run contains the full placeholder
        direct_replacements = 0
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                direct_replacements += 1
        
        if direct_replacements > 0:
            return True
        
        # Method 2: Handle placeholders split across runs
        full_text = ''.join(run.text for run in paragraph.runs)
        
        if old_text not in full_text:
            return False
        
        # Placeholder is split across runs - need to reconstruct
        new_full_text = full_text.replace(old_text, new_text)
        
        # Clear all runs and put the new text in the first run
        # This preserves the paragraph but loses individual run formatting
        for i, run in enumerate(paragraph.runs):
            if i == 0:
                run.text = new_full_text
            else:
                run.text = ''
        
        return True
    
    def replace_in_table_cell(cell):
        """Replace text in all paragraphs within a table cell."""
        replaced = False
        for paragraph in cell.paragraphs:
            if replace_in_paragraph(paragraph):
                replaced = True
        return replaced
    
    # Track replacements
    total_replacements = 0
    
    # Process main document paragraphs
    for paragraph in doc.paragraphs:
        if replace_in_paragraph(paragraph):
            total_replacements += 1
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if replace_in_table_cell(cell):
                    total_replacements += 1
    
    # Process headers and footers
    for section in doc.sections:
        # Headers
        if section.header:
            for paragraph in section.header.paragraphs:
                if replace_in_paragraph(paragraph):
                    total_replacements += 1
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if replace_in_table_cell(cell):
                            total_replacements += 1
        
        # Footers
        if section.footer:
            for paragraph in section.footer.paragraphs:
                if replace_in_paragraph(paragraph):
                    total_replacements += 1
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if replace_in_table_cell(cell):
                            total_replacements += 1
    
    return total_replacements


def find_and_replace_image(doc, placeholder_text, image_path, width=None):
    """Finds a placeholder in a table cell and replaces it with an image. Returns the cell."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder_text in cell.text:
                    cell.text = ''
                    p = cell.paragraphs[0]
                    run = p.add_run()
                    run.add_picture(image_path, width=width)
                    return cell  # Return the cell for caption insertion
    return None


def insert_paragraph_after(paragraph):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = docx.text.paragraph.Paragraph(new_p, paragraph._parent)
    return new_para


def insert_table_after(paragraph, rows, cols):
    """Insert a table after the given paragraph."""
    tbl = OxmlElement('w:tbl')
    
    # Table properties
    tblPr = OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:type'), 'auto')
    tblW.set(qn('w:w'), '0')
    tblPr.append(tblW)
    tbl.append(tblPr)
    
    # Create rows and cells
    for i in range(rows):
        tr = OxmlElement('w:tr')
        for j in range(cols):
            tc = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:type'), 'auto')
            tcPr.append(tcW)
            tc.append(tcPr)
            p = OxmlElement('w:p')
            tc.append(p)
            tr.append(tc)
        tbl.append(tr)
    
    paragraph._element.addnext(tbl)
    return docx.table.Table(tbl, paragraph._parent)


def save_uploaded_file_to_temp(file):
    """Save uploaded file to temporary location and return the file path."""
    if file and file.filename:
        filename = secure_filename(file.filename)
        timestamp = str(int(time.time()))
        filename = f"{timestamp}_{filename}"
        
        # Create temporary file
        temp_dir = tempfile.mkdtemp()
        file_path = os.path.join(temp_dir, filename)
        file.save(file_path)
        return file_path
    return None


def insert_annexure_images(doc, images, captions, placeholder, image_width=Cm(12), image_height=Cm(20), add_final_page_break=True):
    """Insert annexure images one per page at the placeholder location."""
    for para in doc.paragraphs:
        if placeholder in para.text:
            para.text = para.text.replace(placeholder, '')
            insert_after = para
            
            for i, img_path in enumerate(images):
                p_img = insert_paragraph_after(insert_after)
                p_img.paragraph_format.alignment = 1  # Center
                run = p_img.add_run()
                run.add_picture(img_path, width=image_width, height=image_height)
                
                # Insert caption if present
                if captions[i]:
                    p_cap = insert_paragraph_after(p_img)
                    p_cap.add_run(captions[i])
                    p_cap.paragraph_format.alignment = 1  # Center
                    p_cap.runs[0].font.size = Pt(10)
                    insert_after = p_cap
                else:
                    insert_after = p_img
                
                # Page break after each image except the last
                if i < len(images) - 1:
                    p_break = insert_paragraph_after(insert_after)
                    p_break.add_run().add_break(docx.text.run.WD_BREAK.PAGE)
                    insert_after = p_break
            
            # Add final page break if requested
            if add_final_page_break:
                final_break = insert_paragraph_after(insert_after)
                final_break.add_run().add_break(docx.text.run.WD_BREAK.PAGE)
            break  # Only replace the first occurrence


def update_table_of_contents(doc):
    """Update all fields in the document, including table of contents."""
    # Find all field codes in the document
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run._element.xml.find('w:fldChar') != -1:
                # This run contains a field
                run._element.getparent().set(qn('w:dirty'), 'true')
    
    # Find and update TOC fields specifically
    for element in doc.element.body.iter():
        if element.tag.endswith('fldChar'):
            # Mark the field as dirty so it updates when opened
            fld_char_type = element.get(qn('w:fldCharType'))
            if fld_char_type == 'begin':
                # Find the corresponding field instruction
                next_element = element.getnext()
                while next_element is not None:
                    if next_element.tag.endswith('instrText'):
                        if 'TOC' in next_element.text:
                            # Mark TOC field as dirty
                            element.getparent().getparent().set(qn('w:dirty'), 'true')
                        break
                    next_element = next_element.getnext()
