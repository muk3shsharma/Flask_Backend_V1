"""
Image Processing Utilities Module - API Version
==============================================

Specialized utilities for handling gallery and annexure images in Word documents.
Optimized for API usage with temporary file handling.
"""
import os
from docx.shared import Inches, Pt, Cm
import docx
import docx.oxml.shared
from .document_utils import insert_paragraph_after, insert_table_after, save_uploaded_file_to_temp


def insert_gallery_table(doc, images, captions, images_per_row=2, image_width=Cm(8.13), placeholder='{{GALLERY_TABLE}}'):
    """
    Inserts tables at the given placeholder with images and captions.
    Each page displays 6 images (2 per row, 3 rows) with equal alignment.
    Images are sized to 8.13cm width × 5.81cm height.
    """
    for para in doc.paragraphs:
        if placeholder in para.text:
            # Remove placeholder text
            para.text = para.text.replace(placeholder, '')
            
            num_images = len(images)
            if num_images == 0:
                return  # No images to insert
            
            images_per_page = 6  # 2 columns × 3 rows = 6 images per page
            rows_per_page = 3    # Fixed 3 rows per page
            insert_after = para
            
            # Process images in batches of 6 (one page at a time)
            for page_start in range(0, num_images, images_per_page):
                page_end = min(page_start + images_per_page, num_images)
                page_images = images[page_start:page_end]
                page_captions = captions[page_start:page_end]
                
                # Calculate actual rows needed for this page
                page_num_images = len(page_images)
                actual_rows = min(rows_per_page, (page_num_images + images_per_row - 1) // images_per_row)
                
                # Insert table for this page
                table = insert_table_after(insert_after, actual_rows, images_per_row)
                
                # Insert images and captions into the table
                img_idx = 0
                for row in table.rows:
                    for cell in row.cells:
                        if img_idx < page_num_images and page_images[img_idx]:
                            # Clear cell and add image
                            cell.text = ''
                            p_img = cell.paragraphs[0]
                            p_img.alignment = 1  # Center alignment for image
                            run = p_img.add_run()
                            run.add_picture(page_images[img_idx], width=image_width, height=Cm(5.81))
                            
                            # Add caption if exists
                            if page_captions[img_idx]:
                                p_caption = cell.add_paragraph(page_captions[img_idx])
                                p_caption.alignment = 1  # Center alignment
                                p_caption.runs[0].font.size = Pt(10)
                                p_caption.runs[0].font.bold = True
                            
                            # Add cell padding and formatting
                            cell.vertical_alignment = 1  # Center vertical alignment
                            
                            img_idx += 1
                
                # Add page break after each page of images (except the last page)
                if page_end < num_images:
                    # Insert page break after the table
                    table_element = table._element
                    parent = table_element.getparent()
                    new_para_element = docx.oxml.shared.OxmlElement('w:p')
                    new_run_element = docx.oxml.shared.OxmlElement('w:r')
                    new_break_element = docx.oxml.shared.OxmlElement('w:br')
                    new_break_element.set(docx.oxml.shared.qn('w:type'), 'page')
                    new_run_element.append(new_break_element)
                    new_para_element.append(new_run_element)
                    parent.insert(parent.index(table_element) + 1, new_para_element)
                    
                    # Update insert_after to the new paragraph for next page
                    insert_after = table
                else:
                    # For the last page, just add the page break without updating insert_after
                    table_element = table._element
                    parent = table_element.getparent()
                    new_para_element = docx.oxml.shared.OxmlElement('w:p')
                    new_run_element = docx.oxml.shared.OxmlElement('w:r')
                    new_break_element = docx.oxml.shared.OxmlElement('w:br')
                    new_break_element.set(docx.oxml.shared.qn('w:type'), 'page')
                    new_run_element.append(new_break_element)
                    new_para_element.append(new_run_element)
                    parent.insert(parent.index(table_element) + 1, new_para_element)
            
            return


def get_annexure_images_and_captions(prefix, request):
    """Get annexure images and captions from API form data."""
    images = []
    captions = []
    i = 1
    while True:
        img_file = request.files.get(f'{prefix}_image_{i}')
        if not img_file:
            break
        
        img_path = save_uploaded_file_to_temp(img_file)
        if img_path:
            images.append(img_path)
            captions.append(request.form.get(f'{prefix}_caption_{i}', ''))
        
        i += 1
    
    return images, captions


def insert_annexure_images(doc, images, captions, placeholder, image_width=Cm(15), image_height=Cm(20), add_final_page_break=True):
    """
    Inserts each image (with caption) at the given placeholder, one per paragraph, sized to fit within page margins.
    Optionally inserts a page break after the last annexure image.
    Images are centered with portrait proportions: 15cm width × 20cm height.
    """
    from .document_utils import find_and_replace_text
    
    # First, try to find and identify where the placeholder is
    placeholder_found = False
    insert_after_paragraph = None
    
    # Search through all paragraphs
    for para in doc.paragraphs:
        if placeholder in para.text:
            placeholder_found = True
            # Use our improved text replacement function to clear the placeholder
            full_text = ''.join(run.text for run in para.runs)
            if placeholder in full_text:
                # Clear the placeholder properly
                new_text = full_text.replace(placeholder, '')
                # Clear all runs and set new text
                for i, run in enumerate(para.runs):
                    if i == 0:
                        run.text = new_text
                    else:
                        run.text = ''
            
            insert_after_paragraph = para
            break
    
    # If placeholder not found in paragraphs, check tables
    if not placeholder_found:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if placeholder in para.text:
                            placeholder_found = True
                            # Clear placeholder in table cell
                            full_text = ''.join(run.text for run in para.runs)
                            if placeholder in full_text:
                                new_text = full_text.replace(placeholder, '')
                                for i, run in enumerate(para.runs):
                                    if i == 0:
                                        run.text = new_text
                                    else:
                                        run.text = ''
                            insert_after_paragraph = para
                            break
                    if placeholder_found:
                        break
                if placeholder_found:
                    break
            if placeholder_found:
                break
    
    if not placeholder_found:
        print(f"⚠️  Warning: Placeholder {placeholder} not found in document")
        return
    
    if not insert_after_paragraph:
        print(f"⚠️  Warning: Could not determine insertion point for {placeholder}")
        return
    
    print(f"✅ Found {placeholder}, inserting {len(images)} images")
    
    # Insert images after the placeholder location
    current_insert_point = insert_after_paragraph
    
    for i, img_path in enumerate(images):
        try:
            # Insert image with center alignment
            p_img = insert_paragraph_after(current_insert_point)
            p_img.alignment = 1  # Center alignment
            run = p_img.add_run()
            run.add_picture(img_path, width=image_width, height=image_height)
            
            # Insert caption if present
            if i < len(captions) and captions[i]:
                p_cap = insert_paragraph_after(p_img)
                p_cap.add_run(captions[i])
                p_cap.paragraph_format.alignment = 1  # Center
                p_cap.runs[0].font.size = Pt(11)
                p_cap.runs[0].font.bold = True
                # Add some spacing after caption
                p_cap.paragraph_format.space_after = Pt(12)
                current_insert_point = p_cap
            else:
                current_insert_point = p_img
                
            # Page break after each image except the last
            if i < len(images) - 1:
                p_break = insert_paragraph_after(current_insert_point)
                p_break.add_run().add_break(docx.text.run.WD_BREAK.PAGE)
                current_insert_point = p_break
                
            print(f"  ✅ Inserted annexure image {i+1}: {os.path.basename(img_path)}")
            
        except Exception as e:
            print(f"  ❌ Error inserting annexure image {i+1}: {e}")
    
    # Add page break after the last annexure image only if specified
    if add_final_page_break and images:
        try:
            final_break = insert_paragraph_after(current_insert_point)
            final_break.add_run().add_break(docx.text.run.WD_BREAK.PAGE)
            print(f"  ✅ Added final page break after {placeholder}")
        except Exception as e:
            print(f"  ⚠️  Could not add final page break: {e}")
